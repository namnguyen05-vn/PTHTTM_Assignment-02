from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Bao_cao_ASG_02_hoan_chinh.docx"
BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F2F2"
RED = "FCE4D6"
GITHUB_URL = "https://github.com/namnguyen05-vn/PTHTTM_Assignment-02"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def style_run(run, size=10.5, bold=False, color="000000", italic=False):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc, text="", bold_prefix=None, size=10.5, space_after=4, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.08
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        style_run(r1, size=size, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        style_run(r2, size=size)
    else:
        r = p.add_run(text)
        style_run(r, size=size)
    return p


def add_hyperlink(paragraph, text, url, size=10, bold=False):
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    font_size = OxmlElement("w:sz")
    font_size.set(qn("w:val"), str(int(size * 2)))
    run_properties.extend([fonts, color, underline, font_size])
    if bold:
        run_properties.append(OxmlElement("w:b"))
    run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(4 if level == 1 else 3)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    style_run(r, size=17 if level == 1 else 13, bold=True, color="000000")
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.03
    style_run(p.add_run(text), size=10)
    return p


def callout(doc, title, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.45)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 100, 130, 100, 130)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    style_run(p.add_run(title + ": "), size=10, bold=True)
    style_run(p.add_run(text), size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def table(doc, headers, rows, widths=None, font_size=8.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    hdr = t.rows[0]
    set_repeat_table_header(hdr)
    for i, value in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_run(p.add_run(str(value)), size=font_size, bold=True, color="FFFFFF")
        if widths:
            cell.width = Inches(widths[i])
    for row_index, row in enumerate(rows):
        cells = t.add_row().cells
        for i, value in enumerate(row):
            if row_index % 2 == 1:
                set_cell_shading(cells[i], LIGHT_GRAY)
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            style_run(p.add_run(str(value)), size=font_size)
            if widths:
                cells[i].width = Inches(widths[i])
    return t


def add_picture(doc, filename, width=6.3, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    shape = p.add_run().add_picture(str(ROOT / filename), width=Inches(width))
    shape._inline.docPr.set("descr", caption or Path(filename).stem.replace("_", " "))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(4)
        style_run(cp.add_run(caption), size=8.5, italic=True, color="666666")


def add_picture_pair(doc, left_filename, right_filename, width=3.0, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    left_shape = p.add_run().add_picture(str(ROOT / left_filename), width=Inches(width))
    left_shape._inline.docPr.set("descr", Path(left_filename).stem.replace("_", " "))
    p.add_run("   ")
    right_shape = p.add_run().add_picture(str(ROOT / right_filename), width=Inches(width))
    right_shape._inline.docPr.set("descr", Path(right_filename).stem.replace("_", " "))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(4)
        style_run(cp.add_run(caption), size=8.5, italic=True, color="666666")


def add_toc_entry(doc, title, page, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25 if level == 2 else 0)
    p.paragraph_format.right_indent = Inches(0.05)
    p.paragraph_format.space_after = Pt(3 if level == 1 else 1.5)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.1))
    style_run(p.add_run(title), size=10.5 if level == 1 else 9.5, bold=level == 1)
    style_run(p.add_run("\t" + str(page)), size=10.5 if level == 1 else 9.5, bold=level == 1)
    return p


def page_break(doc):
    doc.add_page_break()


def read_csv(filename):
    with open(ROOT / filename, encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def f4(value):
    return f"{float(value):.4f}"


diabetes_val = read_csv("artifacts/diabetes_validation_metrics.csv")
housing_val = read_csv("artifacts/housing_validation_metrics.csv")
ecommerce_val = read_csv("artifacts/ecommerce_validation_metrics.csv")
meta = {
    name: json.loads((ROOT / "models" / f"{name}_metadata.json").read_text(encoding="utf-8"))
    for name in ("diabetes", "housing", "ecommerce")
}


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)
section.header_distance = Inches(0.3)
section.footer_distance = Inches(0.3)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(10.5)
styles["Heading 1"].font.name = "Arial"
styles["Heading 2"].font.name = "Arial"
for footer in [section.footer]:
    add_page_number(footer.paragraphs[0])

# Page 1 — simple Google Docs-style title page.
add_paragraph(doc, "HỌC VIỆN CÔNG NGHỆ BƯU CHÍNH VIỄN THÔNG", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph(doc, "BỘ MÔN PHÁT TRIỂN CÁC HỆ THỐNG THÔNG MINH", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph().paragraph_format.space_after = Pt(70)
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(3)
style_run(title.add_run("PHÁT TRIỂN CÁC HỆ THỐNG THÔNG MINH"), size=26, bold=False)
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(44)
style_run(subtitle.add_run("ASSIGNMENT 02\nFrom Data Representation to a Deployable Intelligent System"), size=14, bold=True)
info = [
    ("Giảng viên hướng dẫn", "TRẦN ĐÌNH QUẾ"),
    ("Họ và tên sinh viên", "NGUYỄN NGỌC HOÀNG NAM"),
    ("Mã sinh viên", "B23DCCN585"),
    ("Lớp", "D23CQCN11-B"),
    ("Nhóm", "05"),
    ("Mã nguồn", "github.com/namnguyen05-vn/PTHTTM_Assignment-02"),
]
table(doc, ["Thông tin", "Nội dung"], info, widths=[2.2, 4.2], font_size=10)
doc.add_paragraph().paragraph_format.space_after = Pt(46)
add_paragraph(doc, "Hà Nội — 2026", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)

# Page 2 — deterministic table of contents for the final layout.
page_break(doc)
heading(doc, "MỤC LỤC", 1)
for title_text, page_number in [
    ("1. Tổng quan và mục tiêu", 3),
    ("2. Data quality, cleaning và chống leakage", 4),
    ("3. Biểu diễn dữ liệu", 5),
    ("4. Ứng dụng 1 — Dự đoán tiểu đường", 6),
    ("5. Ứng dụng 2 — Định giá nhà Hà Nội", 7),
    ("6. Ứng dụng 3 — Khám phá sở thích E-commerce", 8),
    ("7. So sánh chéo và trả lời câu hỏi thảo luận", 10),
    ("8. Web API và mobile-responsive client", 11),
    ("9. Khả năng tái lập, giới hạn và kết luận", 12),
    ("Phụ lục A — Hình ảnh demo Web và Mobile", 13),
]:
    add_toc_entry(doc, title_text, page_number)
add_paragraph(doc, "Mục lục được đối chiếu với bản PDF xuất cuối cùng. Các đề mục con sử dụng hệ thống Heading trong tệp Word để thuận tiện điều hướng và cập nhật khi chỉnh sửa.", size=9, space_after=4)

# Page 3
page_break(doc)
heading(doc, "1. Tổng quan và mục tiêu", 1)
add_paragraph(doc, "Báo cáo trình bày ba hệ thống thông minh theo cùng chuỗi Data → Understand → Clean → Represent → Learn → Evaluate → Persist → Deploy. Trọng tâm không chỉ là độ chính xác mà là tính độc lập của test, biểu diễn số rõ ràng và sự nhất quán giữa notebook, pipeline, API và giao diện.")
table(doc, ["Ứng dụng", "Dữ liệu gốc", "Sau làm sạch", "Đặc trưng triển khai", "Target"], [
    ["Diabetes", "100,000 × 16", "95,655", "6", "diabetes (0/1)"],
    ["Housing", "82,497 × 13", "63,021", "6", "Price (tỷ VNĐ)"],
    ["E-commerce", "23,486 × 11", "22,623", "1 text + 4 số", "Department Name"],
], widths=[1.1, 1.15, 1.1, 1.25, 1.8])
heading(doc, "1.1 Kiến trúc triển khai", 2)
table(doc, ["Dữ liệu", "Làm sạch", "Biểu diễn", "Mô hình", "Triển khai"], [[
    "CSV + review", "Quy tắc + de-dup", "Impute / scale / one-hot / TF-IDF", "Pipeline đã lưu", "FastAPI → responsive web"
]], widths=[.8, 1.2, 1.7, 1.3, 1.55], font_size=8.8)
callout(doc, "Nguyên tắc đánh giá", "So sánh mô hình trên validation; test không tham gia fit preprocessing, huấn luyện, chọn mô hình hoặc chọn ngưỡng. Sau lựa chọn, pipeline được refit bằng train+validation và test đúng một lần.")
heading(doc, "1.2 Nguồn dữ liệu", 2)
for text in [
    "Diabetes Prediction Dataset — kaggle.com/datasets/priyamchoksi/100000-diabetes-clinical-dataset/data",
    "Vietnam Housing Dataset — kaggle.com/code/kerneler/starter-vietnam-housing-dataset-5742ed64-7/input",
    "Women's E-Commerce Clothing Reviews — kaggle.com/datasets/nicapotato/womens-ecommerce-clothing-reviews",
]:
    bullet(doc, text)
heading(doc, "1.3 Định nghĩa bài toán và một quan sát", 2)
for text in [
    "Diabetes: X gồm 6 đặc trưng người bệnh, y là diabetes ∈ {0,1}; một quan sát là một hồ sơ lâm sàng.",
    "Housing: X gồm 6 thuộc tính tin đăng, y là Price theo tỷ VNĐ; một quan sát là một bất động sản được đăng bán.",
    "E-commerce: X gồm Title + Review Text và 4 đặc trưng tabular, y là một trong 6 Department Name; một quan sát là một review cho sản phẩm thời trang.",
]:
    bullet(doc, text)

# Page 3
page_break(doc)
heading(doc, "2. Data quality, cleaning và chống leakage", 1)
table(doc, ["Vấn đề", "Diabetes", "Housing", "E-commerce"], [
    ["Kiểu feature triển khai", "4 numerical + 2 categorical", "3 numerical + 3 categorical", "1 text + 4 numerical"],
    ["Missing", "Không có ở cột chọn", "Nhiều ở số tầng/pháp lý", "845 review; 14 target; Title được phép thiếu"],
    ["Duplicate", "4,345 sau chọn feature", "18,492 bản ghi giống hệt sau parse", "5 comment trùng sau chuẩn hóa"],
    ["Invalid/outlier", "Kiểm soát range tại API", "Đơn vị VN + ngưỡng nghiệp vụ", "Review <2 từ bị loại; API kiểm soát range"],
    ["Class balance", "Lớp bệnh ≈ 8.85%", "Không áp dụng", "Tops 44.4%; Trend 0.52%"],
], widths=[1.15, 1.6, 1.85, 1.75])
heading(doc, "2.1 Thứ tự preprocessing đúng", 2)
for item in [
    "Loại cột ID/không triển khai, sửa kiểu dữ liệu và loại bản ghi trùng trước split.",
    "Tạo train/validation/test độc lập xấp xỉ 70/15/15; classification duy trì phân bố lớp.",
    "Housing chia theo profile X; e-commerce dùng StratifiedGroupKFold theo Clothing ID; overlap group = 0.",
    "Imputer, scaler, one-hot encoder và TF-IDF chỉ fit trong Pipeline trên train.",
    "Chọn model bằng validation; refit train+validation; test một lần; lưu nguyên Pipeline.",
]:
    bullet(doc, item)
heading(doc, "2.2 Cơ chế kiểm soát rò rỉ dữ liệu", 2)
table(doc, ["Nguồn rủi ro", "Biện pháp kiểm soát trong hệ thống"], [
    ["Bản ghi trùng giữa các tập", "De-duplication được thực hiện trước khi chia train/validation/test"],
    ["Thống kê từ validation/test đi vào preprocessing", "Imputer, scaler, encoder và TF-IDF được fit bên trong Pipeline chỉ với train"],
    ["Test ảnh hưởng lựa chọn mô hình", "Mô hình được xếp hạng bằng validation; test chỉ đánh giá pipeline cuối"],
    ["Đặc trưng chứa thông tin target", "Housing loại Giá/m2; e-commerce loại Division Name/Class Name và chỉ dùng Clothing ID để chia group"],
    ["Khác biệt giữa train và inference", "API nạp trực tiếp ba pipeline đã lưu và nhận đúng raw feature schema"],
], widths=[2.35, 4.15])

# Page 4
page_break(doc)
heading(doc, "3. Biểu diễn dữ liệu", 1)
add_paragraph(doc, "Với một quan sát i, dữ liệu tabular sau biến đổi là vector xi ∈ Rᵈ và một batch B quan sát tạo ma trận X ∈ Rᴮˣᵈ. Với text, T là số token và d là số chiều embedding; batch token embedding có dạng E ∈ Rᴮˣᵀˣᵈ.")
table(doc, ["Ứng dụng", "Raw form", "Numerical representation", "Model input"], [
    ["Diabetes", "CSV / table: 4 số + 2 category", "Median impute + scale + one-hot", "B × 13"],
    ["Housing", "CSV / table: 3 số + 3 category", "Median impute + scale + one-hot", "B × 42"],
    ["E-commerce", "CSV + Title/Review Text + 4 số", "TF-IDF + median impute + MaxAbs scale", f"B × {meta['ecommerce']['final_input_shape'][1]}"],
], widths=[1.1, 2.05, 2.25, 1.1])
heading(doc, "3.1 Ý nghĩa kích thước", 2)
for item in [
    "B là số hồ sơ được gửi trong một batch; API hiện dùng B = 1 cho mỗi request.",
    "D là tổng số cột sau biến đổi: feature tabular và vocabulary TF-IDF; e-commerce có D = 8,004.",
    "Trong minh họa text e-commerce: B = 4 review, T = 32 vị trí token và d = 32 chiều embedding.",
    "OneHotEncoder(handle_unknown='ignore') giữ cấu trúc vector ổn định khi inference.",
    "TF-IDF giữ tần suất từ/cụm hai từ nhưng không giữ đầy đủ trật tự dài hạn; embedding SVD nén quan hệ đồng xuất hiện vào 32 chiều.",
]:
    bullet(doc, item)
heading(doc, "3.2 Chuỗi biểu diễn comment e-commerce", 2)
callout(doc, "Comment → Tokens → IDs → Embeddings", "Vocabulary chỉ được học từ train. PAD có ID 0, UNK có ID 1; các token còn lại nhận ID theo vocabulary. Batch minh họa có input_ids [4,32] kiểu int64 và token_embeddings [4,32,32] kiểu float32. Model triển khai dùng TF-IDF + tabular [B,8004].")
heading(doc, "3.3 Training-serving consistency", 2)
add_paragraph(doc, "API tạo DataFrame từ raw JSON với đúng tên cột huấn luyện rồi gọi pipeline.predict/predict_proba. Không có pd.get_dummies, reindex hay fit transformer ở backend. Nhờ vậy cùng một logic preprocessing được dùng ở notebook và production.")
heading(doc, "3.4 Ví dụ từ bản ghi thô đến model input", 2)
table(doc, ["Ứng dụng", "Ví dụ bản ghi thô", "Biểu diễn cuối"], [
    ["Diabetes", "45 tuổi; BMI 28.5; HbA1c 6.5; glucose 150; Male; never", "4 giá trị chuẩn hóa + one-hot = vector 13 chiều"],
    ["Housing", "50 m²; 2 phòng; 3 tầng; Cầu Giấy; nhà ngõ; đã có sổ", "3 giá trị chuẩn hóa + one-hot = vector 42 chiều"],
    ["E-commerce", "'Perfect summer dress'; review; 30 tuổi; rating 5; recommend; feedback 3", "TF-IDF unigram/bigram + 4 số = vector 8,004 chiều"],
], widths=[1.1, 3.15, 2.25], font_size=8)
add_paragraph(doc, "Tham số StandardScaler/MaxAbsScaler và vocabulary TF-IDF đều được học từ train. Các ma trận đầu vào có dtype float64 và dạng B × d.", size=9)

# Page 5 — Diabetes
page_break(doc)
heading(doc, "4. Ứng dụng 1 — Dự đoán tiểu đường", 1)
add_picture(doc, "figures/diabetes_eda.png", width=6.25, caption="Hình 1. Mất cân bằng target và quan hệ HbA1c/glucose với lớp bệnh.")
heading(doc, "4.1 So sánh trên validation", 2)
rows = [[r["Model"], f4(r["Accuracy"]), f4(r["Precision_1"]), f4(r["Recall_1"]), f4(r["F1_1"]), f4(r["ROC_AUC"])] for r in diabetes_val]
table(doc, ["Model", "Accuracy", "P(1)", "R(1)", "F1(1)", "AUC"], rows, widths=[1.6, .9, .8, .8, .8, .8], font_size=7.8)
add_paragraph(doc, "Random Forest có validation F1 lớp bệnh cao nhất (0.7812) nên được chọn. Decision Tree có Recall cao nhưng Precision thấp, tạo nhiều cảnh báo giả. Với y tế, Recall lớp 1 là metric nghiệp vụ quan trọng; F1 lớp 1 được dùng để cân bằng bỏ sót và cảnh báo giả.")
heading(doc, "4.2 Kết quả test cuối", 2)
t = meta["diabetes"]["test_metrics"]
table(doc, ["Accuracy", "Precision lớp 1", "Recall lớp 1", "F1 lớp 1", "ROC-AUC"], [[f4(t["Accuracy"]), f4(t["Precision_1"]), f4(t["Recall_1"]), f4(t["F1_1"]), f4(t["ROC_AUC"])]], widths=[1.25]*5)
add_picture(doc, "figures/diabetes_test_evaluation.png", width=5.5, caption="Hình 2. Test confusion matrix và ROC; FN = 368, TP = 902.")

# Page 6 — Housing
page_break(doc)
heading(doc, "5. Ứng dụng 2 — Định giá nhà Hà Nội", 1)
add_picture(doc, "figures/housing_eda.png", width=6.1, caption="Hình 3. Phân bố giá, diện tích, quận và tương quan biến số.")
heading(doc, "5.1 Làm sạch và tạo target", 2)
add_paragraph(doc, "Chuỗi Giá/m2 được quy đổi về triệu VNĐ/m²: ‘triệu’ giữ nguyên, ‘tỷ’ nhân 1,000, ‘đ/m²’ chia 1,000,000. Target Price = Diện tích × đơn giá / 1,000 (tỷ VNĐ). Sau đó Giá/m2 bị loại khỏi X để tránh target leakage. Bộ lọc cố định: diện tích 10–1,000 m², đơn giá 0.1–1,000 triệu/m², Price 0.1–300 tỷ.")
heading(doc, "5.2 Validation và test", 2)
rows = [[r["Model"], f4(r["MAE"]), f4(r["MSE"]), f4(r["RMSE"]), f4(r["R2"])] for r in housing_val]
table(doc, ["Model", "MAE", "MSE", "RMSE", "R²"], rows, widths=[1.8, 1.0, 1.1, 1.0, .9], font_size=8)
t = meta["housing"]["test_metrics"]
add_paragraph(doc, f"Gradient Boosting có validation MAE thấp nhất và được chọn. Test: MAE = {f4(t['MAE'])} tỷ, RMSE = {f4(t['RMSE'])} tỷ, R² = {f4(t['R2'])}. MAE là sai lệch tiền trung bình; RMSE lớn hơn cho thấy một số căn có sai số lớn.")
callout(doc, "Khoảng bất định", f"API trả prediction ± {meta['housing']['validation_absolute_error_q90']:.3f} tỷ, là bách phân vị 90% của |residual| trên validation. Đây là heuristic vận hành, không phải confidence interval thống kê.")

# Page 7 — E-commerce
page_break(doc)
heading(doc, "6. Ứng dụng 3 — Khám phá sở thích E-commerce", 1)
add_picture(doc, "figures/ecommerce_eda.png", width=5.9, caption="Hình 4. Tần suất nhóm quan tâm, độ dài/rating của review và các từ xuất hiện thường xuyên.")
heading(doc, "6.1 So sánh 6 mô hình", 2)
rows = [[r["Model"], f4(r["Accuracy"]), f4(r["Precision_macro"]), f4(r["Recall_macro"]), f4(r["F1_macro"]), f4(r["ROC_AUC_OVR"])] for r in ecommerce_val]
table(doc, ["Model", "Acc", "P-macro", "R-macro", "F1-macro", "AUC-OVR"], rows, widths=[1.6, .75, .9, .9, .95, .95], font_size=7.5)
add_paragraph(doc, "Logistic Regression đạt validation macro-F1 cao nhất (0.6007) và được chọn theo tiêu chí đã xác định trước. Calibrated Linear SVM có Accuracy cao hơn (0.8309) nhưng macro-F1 thấp hơn (0.5906), cho thấy Accuracy ưu tiên các lớp lớn hơn.")
heading(doc, "6.2 Đóng góp của comments", 2)
ablation = meta["ecommerce"]["representation_ablation_validation"]
table(doc, ["Representation", "Accuracy", "Macro-F1", "AUC-OVR"], [
    [row["Representation"], f4(row["Accuracy"]), f4(row["F1_macro"]), f4(row["ROC_AUC_OVR"])]
    for row in ablation
], widths=[2.8, 1.0, 1.0, 1.0], font_size=8)
add_paragraph(doc, f"Với cùng Logistic Regression, tabular-only đạt macro-F1 {f4(ablation[0]['F1_macro'])}; thêm review TF-IDF đạt {f4(ablation[1]['F1_macro'])}, tăng {float(ablation[1]['F1_macro'])-float(ablation[0]['F1_macro']):.4f}. Như vậy text cung cấp tín hiệu chính để phát hiện nhóm quan tâm.")
heading(doc, "6.3 Kết quả test", 2)
t = meta["ecommerce"]["test_metrics"]
table(doc, ["Accuracy", "P-macro", "R-macro", "F1-macro", "F1-weighted", "AUC-OVR"], [[
    f4(t["Accuracy"]), f4(t["Precision_macro"]), f4(t["Recall_macro"]),
    f4(t["F1_macro"]), f4(t["F1_weighted"]), f4(t["ROC_AUC_OVR"])
]], widths=[1.05]*6, font_size=7.5)
labels = meta["ecommerce"]["classes"]
cm = meta["ecommerce"]["test_confusion_matrix"]
largest_error = max(
    (cm[i][j], labels[i], labels[j])
    for i in range(len(labels)) for j in range(len(labels)) if i != j
)
heading(doc, "6.4 Confusion matrix và phân tích theo lớp", 2)
add_picture(doc, "figures/ecommerce_test_evaluation.png", width=6.15, caption="Hình 5. Confusion matrix và ROC one-vs-rest trên test độc lập.")
per_class = meta["ecommerce"]["test_per_class"]
table(doc, ["Class", "Precision", "Recall", "F1", "Support"], [
    [label, f4(per_class[label]["precision"]), f4(per_class[label]["recall"]), f4(per_class[label]["f1-score"]), int(per_class[label]["support"])]
    for label in labels
], widths=[1.5, 1.2, 1.2, 1.2, 1.0], font_size=7.8)
add_paragraph(doc, f"Accuracy = {f4(t['Accuracy'])} nhưng macro-F1 = {f4(t['F1_macro'])}. Trend chỉ có {int(per_class['Trend']['support'])} mẫu test và recall bằng {f4(per_class['Trend']['recall'])}; vì vậy Accuracy không phản ánh đầy đủ lớp hiếm. Nhầm lẫn lớn nhất là {largest_error[1]} → {largest_error[2]} ({largest_error[0]} review). ROC-AUC OVR = {f4(t['ROC_AUC_OVR'])}.", size=9)

# Page 8 — cross application
page_break(doc)
heading(doc, "7. So sánh chéo và trả lời câu hỏi thảo luận", 1)
table(doc, ["Câu hỏi", "Diabetes", "Housing", "E-commerce"], [
    ["Một observation", "Hồ sơ bệnh nhân", "Tin đăng nhà", "Một review sản phẩm"],
    ["Target", "Nhị phân 0/1", "Giá tỷ VNĐ", "6 Department Name"],
    ["Encode", "Gender, smoking", "Quận, loại nhà, pháp lý", "Token IDs / TF-IDF"],
    ["Normalize", "4 biến số", "3 biến số", "4 biến số MaxAbs"],
    ["Mất thông tin", "Cột không có trên UI", "Địa chỉ/ngày/dài/rộng", "Trật tự dài hạn; hierarchy bị loại"],
    ["Metric chính", "F1 và Recall lớp 1", "MAE", "Macro-F1"],
    ["Model tốt nhất", "Random Forest", "Gradient Boosting", "Logistic Regression"],
    ["Persist", "Pipeline joblib", "Pipeline joblib", "Pipeline joblib"],
], widths=[1.2, 1.7, 1.75, 1.75], font_size=7.8)
heading(doc, "7.1 Metric và quyết định model", 2)
for item in [
    "Diabetes: Recall quan trọng do false negative gây bỏ sót, nhưng chọn bằng F1 lớp 1 để tránh đánh đổi Precision quá cực đoan.",
    "Housing: MAE dễ quy đổi sang tiền; RMSE dùng để nhìn mức phạt các sai số lớn; R² chỉ là tỷ lệ phương sai được giải thích.",
    "E-commerce: macro-F1 cho sáu nhóm trọng số ngang nhau và làm lộ khó khăn ở lớp Trend; weighted-F1 và Accuracy vẫn được báo để có bối cảnh.",
]:
    bullet(doc, item)
heading(doc, "7.2 Thông tin được giữ và mất", 2)
add_paragraph(doc, "One-hot giữ membership của category nhưng mất quan hệ thứ bậc ngầm định. Scaling giữ tương quan tuyến tính nhưng thay đổi đơn vị. Với e-commerce, TF-IDF giữ từ/cụm từ nổi bật nhưng không mô hình hóa đầy đủ ngữ cảnh xa; loại Division Name và Class Name làm giảm thông tin song ngăn target leakage.")

# Page 9 — deployment
page_break(doc)
heading(doc, "8. Web API và mobile-responsive client", 1)
table(doc, ["Endpoint", "Input", "Output chính"], [
    ["POST /predict/diabetes", "6 raw features", "class, risk probability, confidence"],
    ["POST /predict/housing", "6 raw features", "price + heuristic range"],
    ["POST /predict/ecommerce", "review/title + 4 raw features", "interest + 6 probabilities"],
    ["GET /health", "—", "loaded pipelines"],
], widths=[1.8, 1.7, 3.0])
heading(doc, "8.1 Tính nhất quán và an toàn đầu vào", 2)
for item in [
    "Pydantic cấm field thừa, kiểm tra range số và category hợp lệ.",
    "Đường dẫn dùng Path(__file__), không phụ thuộc working directory.",
    "Frontend dùng window.location.origin, không hard-code IP nội bộ.",
    "CORS mặc định chỉ cho localhost; có thể cấu hình ALLOWED_ORIGINS.",
    "HTML escape dữ liệu bảng trước khi render để giảm rủi ro chèn nội dung.",
]:
    bullet(doc, item)
heading(doc, "8.2 Minh họa responsive", 2)
add_picture(doc, "figures/deployment_architecture.png", width=6.25, caption="Hình 6. Responsive web và REST API dùng trực tiếp ba pipeline đã persist.")
add_paragraph(doc, "Giao diện Bootstrap Grid co về một cột trên điện thoại và thực hiện đầy đủ luồng mobile client: nhập dữ liệu → gửi REST request → nhận kết quả → hiển thị prediction/confidence. Cùng một giao diện phục vụ trình duyệt desktop, điện thoại và máy tính bảng.", size=9.5)
callout(doc, "Bằng chứng chạy", "tests/test_api.py kiểm tra health, trang web, ba endpoint dự đoán và lỗi validation; kết quả hiện tại: 3 tests passed.")

# Page 10 — reproducibility and conclusion
page_break(doc)
heading(doc, "9. Khả năng tái lập, giới hạn và kết luận", 1)
heading(doc, "9.1 Reproducibility", 2)
table(doc, ["Thành phần", "Bằng chứng"], [
    ["Notebook", "3 file .ipynb đã chạy, không có error output"],
    ["Model", "3 file *_pipeline.joblib + 3 metadata JSON"],
    ["Metric", "CSV validation + confusion matrix/ROC hoặc residual plot"],
    ["Môi trường", "requirements.txt pin phiên bản; random_state=42"],
    ["Triển khai", "main.py + index.html + tests/test_api.py"],
], widths=[1.6, 4.9])
github_paragraph = doc.add_paragraph()
github_paragraph.paragraph_format.space_after = Pt(5)
style_run(github_paragraph.add_run("Mã nguồn và hướng dẫn chạy: "), size=10, bold=True)
add_hyperlink(github_paragraph, "GitHub repository PTHTTM_Assignment-02", GITHUB_URL, size=10)
heading(doc, "9.2 Giới hạn và hướng phát triển", 2)
table(doc, ["Giới hạn", "Ảnh hưởng", "Hướng phát triển"], [
    ["Trend chỉ có 118 review sạch", "Macro-F1 và recall lớp hiếm thấp", "Thu thập thêm review Trend hoặc thử cost-sensitive learning"],
    ["Dữ liệu chỉ thuộc thời trang nữ", "Không tổng quát sang ngành hàng khác", "Đánh giá ngoài mẫu trên catalog e-commerce khác"],
    ["TF-IDF hạn chế ngữ cảnh dài", "Có thể bỏ sót phủ định/ngữ nghĩa tinh tế", "So sánh thêm contextual embeddings khi có tài nguyên"],
    ["Housing R² ≈ 0.52", "Chưa giải thích hết biến động giá", "Thêm tọa độ, thời gian đăng, mặt đường và khoảng cách trung tâm"],
    ["Diabetes dùng 6 feature triển khai", "Một phần thông tin lâm sàng không được dùng", "Mở rộng đồng thời form, schema API và pipeline"],
], widths=[1.8, 2.1, 2.6], font_size=8)
heading(doc, "9.3 Kết luận", 2)
add_paragraph(doc, "Ba ứng dụng tuân thủ cùng một vòng đời ML có thể kiểm chứng: làm sạch trước split, transformer fit trong pipeline, lựa chọn bằng validation, đánh giá test độc lập, lưu pipeline và inference bằng đúng raw schema. Quy trình này tạo ra kết quả tái lập, giữ nhất quán giữa notebook và dịch vụ triển khai, đồng thời xác định rõ phạm vi áp dụng của từng mô hình.")
heading(doc, "Tài liệu tham chiếu", 2)
for item in [
    "Scikit-learn User Guide — Pipeline, ColumnTransformer, model evaluation.",
    "FastAPI documentation — request validation, OpenAPI and deployment.",
    "Đề bài Assignment 02 và Lecture 02 — Data Representation, PTIT (2026).",
]:
    bullet(doc, item)

# Pages 13–16 — screenshots supplied by the author, limited to flows that
# match the current source code and persisted models.
page_break(doc)
heading(doc, "Phụ lục A — Hình ảnh demo Web và Mobile", 1)
add_paragraph(doc, "Phụ lục minh họa các luồng nhập dữ liệu và trả kết quả trên trình duyệt desktop và giao diện responsive trên điện thoại. Các ảnh được lựa chọn từ đúng phiên bản source code hiện tại.", size=10)
heading(doc, "A.1 Web — Dự đoán tiểu đường", 2)
add_picture(doc, "demo/Demo-Web_Diabetes_Input.png", width=5.75, caption="Hình 7. Form nhập liệu tiểu đường trên trình duyệt desktop.")
add_picture(doc, "demo/Demo-Web_Diabetes_Output.png", width=5.75, caption="Hình 8. Kết quả dự đoán, độ tin cậy và các trường hợp tương tự.")

page_break(doc)
heading(doc, "A.2 Web — Định giá bất động sản", 2)
add_picture(doc, "demo/Demo-Web_Housing_Price_Input.png", width=6.0, caption="Hình 9. Form định giá bất động sản trên trình duyệt desktop.")
add_picture(doc, "demo/Demo-Web_Housing_Price_Output.png", width=5.75, caption="Hình 10. Giá ước tính, khoảng dao động heuristic và các trường hợp tương tự.")

page_break(doc)
heading(doc, "A.3 Mobile-responsive — Dự đoán tiểu đường", 2)
add_picture_pair(
    doc,
    "demo/Demo-Mobile_Diabetes_Input.jpg",
    "demo/Demo-Mobile_Diabetes_Output.jpg",
    width=2.65,
    caption="Hình 11–12. Luồng nhập dữ liệu (trái) và kết quả (phải) trên điện thoại.",
)

page_break(doc)
heading(doc, "A.4 Mobile-responsive — Định giá bất động sản", 2)
add_picture_pair(
    doc,
    "demo/Demo-Mobile_Housing_Price_Input.jpg",
    "demo/Demo-Mobile_Housing_Price_Output.jpg",
    width=2.65,
    caption="Hình 13–14. Luồng nhập dữ liệu (trái) và kết quả định giá (phải) trên điện thoại.",
)

doc.core_properties.title = "Assignment 02 — Intelligent Systems Development"
doc.core_properties.author = "Nguyễn Ngọc Hoàng Nam"
doc.core_properties.subject = "Technical report synchronized with notebooks, persisted pipelines, API and responsive web client"
doc.save(OUT)
print(OUT)
